import os
import re
import hashlib
import json
import math
import shutil
import pandas as pd
import pymupdf4llm
from docx import Document
from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document as LangDocument
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownTextSplitter, 
    MarkdownHeaderTextSplitter,)

SECTION_ALIASES = {
    "abstract": "ABSTRACT",
    "introduction": "INTRODUCTION",
    "background": "INTRODUCTION",
    "method": "METHODS",
    "methods": "METHODS",
    "methodology": "METHODS",
    "materials and methods": "METHODS",
    "results": "RESULTS",
    "discussion": "DISCUSSION",
    "conclusion": "CONCLUSION",
    "conclusions": "CONCLUSION",
    "references": "REFERENCES",
}

SECTION_REGEX = re.compile(
    r'^\s*(abstract|introduction|background|methodology|methods?|materials and methods|results?|discussion|conclusions?|references)\s*$',
    re.IGNORECASE
)

PDF_CHUNK_SIZE = int(os.getenv("PDF_CHUNK_SIZE", "1400"))
PDF_CHUNK_OVERLAP = int(os.getenv("PDF_CHUNK_OVERLAP", "120"))
TEXT_CHUNK_SIZE = int(os.getenv("TEXT_CHUNK_SIZE", "1100"))
TEXT_CHUNK_OVERLAP = int(os.getenv("TEXT_CHUNK_OVERLAP", "120"))
PDF_CACHE_DIR = os.getenv("PDF_CACHE_DIR", os.path.join("uploads", ".pdf_cache"))
PDF_TABLE_FACT_LIMIT = int(os.getenv("PDF_TABLE_FACT_LIMIT", "24"))


def _format_summary_value(value) -> str:
    if pd.isna(value):
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def _clean_tabular_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.dropna(how="all").copy()
    cleaned.columns = [str(c).strip().lower() for c in cleaned.columns]
    return cleaned


def _looks_numeric(series: pd.Series) -> bool:
    numeric = pd.to_numeric(series, errors="coerce")
    return numeric.notna().sum() >= max(3, int(len(series) * 0.6))


def _looks_datetime(series: pd.Series) -> bool:
    converted = pd.to_datetime(series, errors="coerce")
    return converted.notna().sum() >= max(3, int(len(series) * 0.6))


def _format_number(value) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return _format_summary_value(value)
    if math.isfinite(number) and number.is_integer():
        return str(int(number))
    return f"{number:.2f}".rstrip("0").rstrip(".")


def _humanize_column_name(name: str) -> str:
    text = str(name).strip().replace("_", " ")
    return re.sub(r"\s+", " ", text)


def _is_numeric_value(value) -> bool:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return False
    return math.isfinite(number)


def _build_dataset_summary_text(df: pd.DataFrame) -> str:
    lines = [
        "Tabular dataset summary",
        f"Total rows: {len(df)}",
        f"Columns: {', '.join(str(c) for c in df.columns)}",
    ]

    numeric_cols = []
    categorical_cols = []
    datetime_cols = []

    for col in df.columns:
        series = df[col]
        if _looks_numeric(series):
            numeric_cols.append(col)
        elif _looks_datetime(series):
            datetime_cols.append(col)
        else:
            categorical_cols.append(col)

    if numeric_cols:
        lines.append("Numeric stats:")
        for col in numeric_cols[:12]:
            numeric = pd.to_numeric(df[col], errors="coerce").dropna()
            if numeric.empty:
                continue
            lines.append(
                f"{col}: min {_format_number(numeric.min())}, max {_format_number(numeric.max())}, avg {_format_number(numeric.mean())}"
            )

    if datetime_cols:
        lines.append("Date coverage:")
        for col in datetime_cols[:8]:
            parsed = pd.to_datetime(df[col], errors="coerce").dropna()
            if parsed.empty:
                continue
            lines.append(
                f"{col}: {parsed.min().date().isoformat()} to {parsed.max().date().isoformat()}"
            )

    if categorical_cols:
        lines.append("Categorical columns:")
        for col in categorical_cols[:12]:
            cleaned = df[col].map(_format_summary_value)
            values = [v for v in cleaned.tolist() if v]
            if not values:
                continue
            unique_values = list(dict.fromkeys(values))
            preview = ", ".join(unique_values[:6])
            if len(unique_values) > 6:
                preview += f", +{len(unique_values) - 6} more"
            lines.append(f"{col}: {preview}")

    return normalize_text("\n".join(lines))


def build_chunk_text(df: pd.DataFrame) -> str:
    summary = []

    if "sales" in df.columns:
        sales = pd.to_numeric(df["sales"], errors="coerce").dropna()
        if not sales.empty:
            summary.append(f"Total sales: {_format_number(sales.sum())}")
            summary.append(f"Average sales: {sales.mean():.2f}")

    if "country" in df.columns:
        countries = [
            value
            for value in dict.fromkeys(df["country"].dropna().astype(str).str.strip())
            if value
        ]
        if countries:
            summary.append(f"Countries: {', '.join(countries)}")

    if not summary:
        summary.append(f"Total rows: {len(df)}")
        summary.append(f"Columns: {', '.join(str(col) for col in df.columns)}")

    return normalize_text("\n".join(summary))


def add_schema_context(df: pd.DataFrame) -> str:
    lines = ["This dataset contains:"]
    for col in df.columns[:12]:
        label = _humanize_column_name(col)
        description = "descriptive field"
        series = df[col]
        lower_col = str(col).lower()
        if _looks_numeric(series):
            description = "numeric business metric"
        elif _looks_datetime(series) or "date" in lower_col or "time" in lower_col:
            description = "date or time field"
        elif lower_col in {"country", "region", "location"}:
            description = "location of sales or business activity"
        elif lower_col in {"product", "item", "sku"}:
            description = "product or item being sold"
        elif lower_col in {"sales", "revenue", "amount"}:
            description = "sales or revenue value"
        lines.append(f"- {label}: {description}")
    return normalize_text("\n".join(lines))


def enhance_chunk(text: str) -> str:
    prefix = "This data describes sales performance, revenue, transactions, products, countries, and business trends."
    return normalize_text(f"{prefix}\n{text}")


def add_global_summary(df: pd.DataFrame) -> str:
    sales_col = "sales" if "sales" in df.columns else None
    if sales_col is None:
        return ""

    sales = pd.to_numeric(df[sales_col], errors="coerce").dropna()
    if sales.empty:
        return ""

    return normalize_text(
        "\n".join(
            [
                "Global Summary:",
                f"Total sales: {_format_number(sales.sum())}",
                f"Average sales: {_format_number(sales.mean())}",
                f"Max sales: {_format_number(sales.max())}",
                f"Min sales: {_format_number(sales.min())}",
            ]
        )
    )


def _extract_representative_metadata(df: pd.DataFrame) -> dict:
    metadata = {}

    if "country" in df.columns:
        countries = [v for v in df["country"].dropna().astype(str).str.strip() if v]
        if countries:
            metadata["country"] = countries[0]
            metadata["countries"] = list(dict.fromkeys(countries[:50]))

    if "product" in df.columns:
        products = [v for v in df["product"].dropna().astype(str).str.strip() if v]
        if products:
            metadata["product"] = products[0]
            metadata["products"] = list(dict.fromkeys(products[:50]))

    return metadata


def _row_to_semantic_sentence(row: pd.Series, columns: list[str]) -> str:
    parts = []
    subject = None

    preferred_subjects = ("product", "company", "customer", "name", "item", "country", "region")
    for col in columns:
        value = _format_summary_value(row[col])
        if not value:
            continue
        if subject is None and any(token in col for token in preferred_subjects):
            subject = value
            break

    for col in columns:
        value = _format_summary_value(row[col])
        if not value:
            continue
        label = _humanize_column_name(col)
        lower_col = col.lower()
        if _is_numeric_value(row[col]):
            parts.append(f"{label} is {_format_number(row[col])}")
        elif "date" in lower_col or "time" in lower_col:
            parts.append(f"{label} is {value}")
        else:
            parts.append(f"{label} is {value}")

    if not parts:
        return ""
    if subject:
        return normalize_text(f"{subject}. " + ". ".join(parts) + ".")
    return normalize_text(". ".join(parts) + ".")


def row_to_narrative(row: pd.Series, columns: list[str]) -> str:
    values = {col: _format_summary_value(row[col]) for col in columns}
    country = values.get("country", "")
    product = values.get("product", values.get("item", ""))
    sales = values.get("sales", values.get("revenue", values.get("amount", "")))
    date = values.get("date", values.get("order date", values.get("order_date", "")))

    narrative_parts = []
    if country and product and sales:
        sentence = f"In {country}, product {product} generated sales of {sales}"
        if date:
            sentence += f" on {date}"
        narrative_parts.append(sentence + ".")
    elif product and sales:
        sentence = f"Product {product} generated sales of {sales}"
        if date:
            sentence += f" on {date}"
        narrative_parts.append(sentence + ".")

    semantic_sentence = _row_to_semantic_sentence(row, columns)
    if semantic_sentence:
        narrative_parts.append(semantic_sentence)

    return normalize_text(" ".join(narrative_parts))


def _build_block_text(part: pd.DataFrame, *, block_label: str) -> str:
    """Build retrieval-friendly text for a grouped set of tabular rows."""
    lines = [
        block_label,
        f"Rows in block: {len(part)}",
        add_schema_context(part),
    ]
    chunk_summary = build_chunk_text(part)
    if chunk_summary:
        lines.append(chunk_summary)

    for col in part.columns:
        label = str(col).strip() or "Column"
        series = part[col]
        cleaned = series.map(_format_summary_value)
        non_empty = cleaned[cleaned != ""]
        if non_empty.empty:
            continue

        if _looks_numeric(series):
            numeric = pd.to_numeric(series, errors="coerce").dropna()
            if numeric.empty:
                continue
            low = numeric.min()
            high = numeric.max()
            low_text = _format_number(low)
            high_text = _format_number(high)
            if low == high:
                lines.append(f"{label}: {low_text}")
            else:
                lines.append(f"{label} range: {low_text}-{high_text}")
                lines.append(f"Average {label}: {_format_number(numeric.mean())}")
            continue

        unique_values = list(dict.fromkeys(non_empty.tolist()))
        if len(unique_values) == 1:
            lines.append(f"{label}: {unique_values[0]}")
            continue

        preview = ", ".join(unique_values[:6])
        if len(unique_values) > 6:
            preview += f", +{len(unique_values) - 6} more"
        lines.append(f"{label}: {preview}")

    narratives = []
    columns = [str(c) for c in part.columns]
    for _, row in part.head(8).iterrows():
        narrative = row_to_narrative(row, columns)
        if narrative:
            narratives.append(narrative)

    if narratives:
        lines.append("Example transactions:")
        lines.extend(narratives)

    return enhance_chunk("\n".join(lines))


def ensure_pdf_cache_dir() -> str:
    os.makedirs(PDF_CACHE_DIR, exist_ok=True)
    return PDF_CACHE_DIR


def file_sha256(file_path: str, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with open(file_path, "rb") as file_obj:
        while True:
            chunk = file_obj.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def pdf_cache_path(file_path: str) -> str:
    cache_dir = ensure_pdf_cache_dir()
    file_hash = file_sha256(file_path)
    base_name = os.path.basename(file_path)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", base_name)
    return os.path.join(cache_dir, f"{safe_name}.{file_hash}.json")


def load_cached_pdf_docs(file_path: str) -> list[LangDocument] | None:
    cache_path = pdf_cache_path(file_path)
    if not os.path.exists(cache_path):
        return None

    with open(cache_path, "r", encoding="utf-8") as file_obj:
        cached_pages = json.load(file_obj)

    docs = []
    for page in cached_pages:
        metadata = page.get("metadata") or {"page": page.get("page", -1)}
        text = normalize_text(page.get("text", ""))
        if not text:
            continue
        docs.append(
            LangDocument(
                page_content=text,
                metadata=metadata,
            )
        )
    return docs


def save_cached_pdf_docs(file_path: str, docs: list[LangDocument]) -> None:
    cache_path = pdf_cache_path(file_path)
    cached_pages = [
        {
            "page": doc.metadata.get("page", -1),
            "metadata": doc.metadata,
            "text": doc.page_content,
        }
        for doc in docs
        if getattr(doc, "page_content", None)
    ]
    with open(cache_path, "w", encoding="utf-8") as file_obj:
        json.dump(cached_pages, file_obj, ensure_ascii=True)


def configure_tesseract() -> str | None:
    """Expose a local Tesseract install to libraries that depend on PATH lookup."""
    candidates = []
    env_cmd = os.environ.get("TESSERACT_CMD")
    if env_cmd:
        candidates.append(env_cmd)

    which_cmd = shutil.which("tesseract")
    if which_cmd:
        candidates.append(which_cmd)

    if os.name == "nt":
        candidates.extend(
            [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                os.path.join(
                    os.environ.get("LOCALAPPDATA", ""),
                    "Programs",
                    "Tesseract-OCR",
                    "tesseract.exe",
                ),
            ]
        )

    for cmd in candidates:
        if not cmd or not os.path.exists(cmd):
            continue
        install_dir = os.path.dirname(cmd)
        path_parts = os.environ.get("PATH", "").split(os.pathsep)
        if install_dir not in path_parts:
            os.environ["PATH"] = install_dir + os.pathsep + os.environ.get("PATH", "")
        return cmd

    return None

def inject_section_markers(text: str) -> str:
    """
    convert raw pdf text into rag sections
    required for your  /query section retrieval to work
    """
    lines = text.split("\n")
    new_lines = []
    
    for line in lines:
        clean = line.strip().lower()

        if SECTION_REGEX.match(clean):
            section_name = SECTION_ALIASES.get(clean, clean).upper()
            new_lines.append(f"[SECTION:{section_name}]")

        new_lines.append(line)

    return "\n".join(new_lines)


def _clean_table_cell(text: str) -> str:
    text = normalize_text(str(text or ""))
    text = text.strip("|").strip()
    return text


def extract_pdf_table_facts(text: str, page_no) -> list[LangDocument]:
    """Convert simple markdown table rows into compact fact chunks for retrieval."""
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    facts = []

    for line in lines:
        if "|" not in line:
            continue

        cells = [_clean_table_cell(cell) for cell in line.split("|")]
        cells = [cell for cell in cells if cell]
        if len(cells) < 2:
            continue

        if all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
            continue

        header = cells[0].lower()
        if header in {"metric", "measure", "description", "item", "note"} and len(cells) == 2:
            continue

        key = cells[0]
        values = cells[1:]
        if not any(re.search(r"\d|[$%\u20ac\u00a3\u00a5]", value) for value in values):
            continue

        for value in values:
            fact_text = normalize_text(f"{key}: {value}")
            if not fact_text or len(fact_text) < 4:
                continue
            facts.append(
                LangDocument(
                    page_content=fact_text,
                    metadata={
                        "page": page_no,
                        "table_fact": True,
                    },
                )
            )
            if len(facts) >= PDF_TABLE_FACT_LIMIT:
                return facts

    return facts

def normalize_text(text:str) -> str:
    """normalize whitespace for stable embeddings"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Keep paragraph boundaries for section heading detection.
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.replace("\x00", "")
    return text.strip()

def made_chunk_id(file_path: str, content: str, page=None) -> str:
    """Stable deterministic ID for updates"""
    base = f"{file_path}|{page}|{content[:200]}"
    return hashlib.sha256(base.encode()).hexdigest()[:32]

def attach_metadata(docs: list[LangDocument], file_path: str, ext: str) -> list[LangDocument]:
    base =os.path.basename(file_path)
    for d in docs:
        d.metadata = d.metadata or {}
        d.metadata["file"] = base
        d.metadata["source"] = file_path
        d.metadata["file_type"] = ext
    return docs
def load_pdf(file_path: str) -> list[LangDocument]:
    """Layout aware pdf loader"""
    cached_docs = load_cached_pdf_docs(file_path)
    if cached_docs is not None:
        print(f"[PDF CACHE HIT] {os.path.basename(file_path)} -> {len(cached_docs)} pages")
        return cached_docs

    tesseract_cmd = configure_tesseract()
    try:
        md = pymupdf4llm.to_markdown(file_path, page_chunks=True)
    except Exception as exc:
        message = str(exc)
        if "tesseract" in message.lower() or "ocr" in message.lower():
            if tesseract_cmd:
                raise RuntimeError(
                    f"OCR failed for {os.path.basename(file_path)} even though Tesseract was found at "
                    f"{tesseract_cmd}. Restart the app and try the upload again."
                ) from exc
            raise RuntimeError(
                f"OCR is required for {os.path.basename(file_path)} but Tesseract was not detected. "
                "Install Tesseract or set the TESSERACT_CMD environment variable."
            ) from exc
        raise
    docs = []
    for page in md:
        if not isinstance(page, dict):
            continue

        page_no = page.get("page", page.get("page_number", page.get("number", -1)))
        raw_text = page.get("text", "")
        text = normalize_text(raw_text)
        text = inject_section_markers(text)
        table_fact_docs = extract_pdf_table_facts(raw_text, page_no)

        if text:
            docs.append(
                LangDocument(
                    page_content=text,
                    metadata={"page": page_no}
                )
            )

        docs.extend(table_fact_docs)

        if not text and not table_fact_docs:
            continue
    save_cached_pdf_docs(file_path, docs)
    print(f"[PDF CACHE SAVE] {os.path.basename(file_path)} -> {len(docs)} pages")
    return docs

def load_docx(file_path: str) -> list[LangDocument]:
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [LangDocument(page_content=normalize_text(text))]


def load_csv_excel(file_path: str, ext: str) -> list[LangDocument]:
    """Build summary-first, column-aware semantic chunks for tabular files."""
    df = pd.read_csv(file_path) if ext == "csv" else pd.read_excel(file_path)
    df = _clean_tabular_dataframe(df)

    row_count = len(df)
    if row_count <= 200:
        chunk_rows = 1
    elif row_count <= 2000:
        chunk_rows = 25
    else:
        chunk_rows = 50
    max_blocks = None
    if row_count > 12500:
        max_blocks = 250

    total_blocks = max(1, math.ceil(row_count / chunk_rows)) if row_count else 0
    selected_block_indexes = list(range(total_blocks))
    if max_blocks is not None and total_blocks > max_blocks:
        step = total_blocks / max_blocks
        selected_block_indexes = []
        seen = set()
        for i in range(max_blocks):
            block_idx = min(total_blocks - 1, int(i * step))
            if block_idx in seen:
                continue
            seen.add(block_idx)
            selected_block_indexes.append(block_idx)

        # Keep the tail represented for time-series style sheets.
        if selected_block_indexes and selected_block_indexes[-1] != total_blocks - 1:
            selected_block_indexes[-1] = total_blocks - 1

    blocks = []
    semantic_rows = 0
    semantic_index_capped = max_blocks is not None and total_blocks > max_blocks

    global_summary_text = add_global_summary(df)
    if global_summary_text:
        blocks.append(
            LangDocument(
                page_content=enhance_chunk(
                    "\n".join(
                        [
                            add_schema_context(df),
                            global_summary_text,
                        ]
                    )
                ),
                metadata={
                    **_extract_representative_metadata(df),
                    "rows": "global_summary",
                    "total_rows": row_count,
                    "semantic_rows_indexed": 0,
                    "semantic_blocks_total": total_blocks,
                    "semantic_blocks_indexed": len(selected_block_indexes),
                    "semantic_index_capped": semantic_index_capped,
                    "chunk_rows": chunk_rows,
                    "chunk_kind": "global_summary",
                    "special_chunk": True,
                },
            )
        )

    summary_metadata = {
        **_extract_representative_metadata(df),
        "rows": "summary",
        "total_rows": row_count,
        "semantic_rows_indexed": 0,
        "semantic_blocks_total": total_blocks,
        "semantic_blocks_indexed": len(selected_block_indexes),
        "semantic_index_capped": semantic_index_capped,
        "chunk_rows": chunk_rows,
        "chunk_kind": "dataset_summary",
    }
    blocks.append(
        LangDocument(
            page_content=enhance_chunk(
                "\n".join(
                    filter(None, [add_schema_context(df), _build_dataset_summary_text(df), build_chunk_text(df)])
                )
            ),
            metadata=summary_metadata,
        )
    )

    for block_idx in selected_block_indexes:
        i = block_idx * chunk_rows
        part = df.iloc[i:i+chunk_rows].copy()
        if chunk_rows == 1:
            block_label = f"Row {i}"
        else:
            block_label = f"Data block rows {i}-{i + len(part) - 1}"
        table_text = _build_block_text(part, block_label=block_label)
        semantic_rows += len(part)
        blocks.append(
            LangDocument(
                page_content=table_text,
                metadata={
                    **_extract_representative_metadata(part),
                    "rows": f"{i}-{i+len(part)-1}",
                    "total_rows": row_count,
                    "semantic_rows_indexed": semantic_rows,
                    "semantic_blocks_total": total_blocks,
                    "semantic_blocks_indexed": len(selected_block_indexes),
                    "semantic_index_capped": semantic_index_capped,
                    "chunk_rows": chunk_rows,
                    "chunk_kind": "row" if chunk_rows == 1 else "data_block",
                }
            )
        )

    return blocks

# MAIN ENTRY

def load_and_chunk_file(file_path: str) -> list[LangDocument]:


    ext = os.path.splitext(file_path)[1].lower().replace(".", "")

    # -------- Load --------
    if ext == "pdf":
        docs = load_pdf(file_path)
    elif ext == "txt":
        docs = TextLoader(file_path, encoding="utf-8").load()
    elif ext == "docx":
        docs = load_docx(file_path)
    elif ext in ["csv", "xls", "xlsx"]:
        docs = load_csv_excel(file_path, ext)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    docs = attach_metadata(docs, file_path, ext)

    # CHUNKING STRATEGY
    
    if ext == "pdf":
        splitter = RecursiveCharacterTextSplitter(
            chunk_size = PDF_CHUNK_SIZE,
            chunk_overlap = PDF_CHUNK_OVERLAP,
            separators = [
                "\n[SECTION:",
                "\n\n",
                "\n",
                ". ",
                " "
            ]
        )
        chunks = splitter.split_documents(docs)
    elif ext in ["csv", "xls", "xlsx"]:
        chunks = docs
    else:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=TEXT_CHUNK_SIZE,
            chunk_overlap=TEXT_CHUNK_OVERLAP,
        )
        chunks = splitter.split_documents(docs)

    
    # FINALIZE (IDs + CLEAN)
    
    final_chunks = []
    seen = set()
    for c in chunks:
        content = normalize_text(c.page_content)
        if not content:
            continue
        page = c.metadata.get("page")
        cid = made_chunk_id(file_path, content, page)
        if cid in seen:
            continue
        seen.add(cid)
        c.page_content = content
        c.metadata["chunk_id"] = cid
        final_chunks.append(c)
    print(f"[INGEST] {os.path.basename(file_path)} -> {len(final_chunks)} chunks")
    return final_chunks


def load_file(file_path: str) -> list[LangDocument]:
    """Backward-compatible entrypoint used by app.py."""
    return load_and_chunk_file(file_path)




    
